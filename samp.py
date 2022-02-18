#page 2


import tkinter as tk
import ctypes
from tkinter import filedialog as fd
from PIL import ImageTk, Image

from tkinter import Tk, Canvas, Button
import pytesseract
import docx
import webbrowser

ctypes.windll.shcore.SetProcessDpiAwareness(2,3)
#creating a window

root = tk.Tk()
root.geometry('1300x850+290+50')
root.title('   ')

canvas = Canvas(
    root,
    height = 1300,
    width = 850)


img = ImageTk.PhotoImage(Image.open('D:\pyexpo\pg2 2.jpg'))
panel = tk.Label(root, image = img)
panel.pack(side = "bottom", fill = "both", expand = "yes")

#jpg
def to_jpg():
    filename=fd.askopenfilename()


    if filename.endswith(".png") or filename.endswith(".pdf") or filename.endswith(".gif") or filename.endswith(".tiff") or filename.endswith(".tiff") :
       Image.open(filename).save(fd.asksaveasfilename(defaultextension='.jpg'))

def to_pdf():
    filename=fd.askopenfilename()


    if filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".gif") or filename.endswith(".tiff") or filename.endswith(".tiff") or filename.endswith('.jpeg') :
       Image.open(filename).save(fd.asksaveasfilename(defaultextension='.pdf'))

def to_tiff():
    filename=fd.askopenfilename()


    if filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".gif") or filename.endswith(".tiff") or filename.endswith(".pdf") or filename.endswith('.jpeg') :
       Image.open(filename).save(fd.asksaveasfilename(defaultextension='.tiff'))


def to_gif():
    filename=fd.askopenfilename()


    if filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".gif") or filename.endswith(".tiff") or filename.endswith(".pdf") or filename.endswith('.jpeg') :
       Image.open(filename).save(fd.asksaveasfilename(defaultextension='.gif'))

def to_png():
    filename=fd.askopenfilename()


    if filename.endswith(".png") or filename.endswith(".jpg") or filename.endswith(".gif") or filename.endswith(".tiff") or filename.endswith(".pdf") or filename.endswith('.jpeg') :
       Image.open(filename).save(fd.asksaveasfilename(defaultextension='.png'))

#back
def back():
    root.destroy()
    import sample
#english
def eng():
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    ocr = pytesseract.image_to_string(fd.askopenfilename())
    print(ocr)
    doc = docx.Document()
    doc.add_paragraph(ocr)
    doc.save(fd.asksaveasfilename(defaultextension='.docx'))

#tamil
def tam():
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    ocr = pytesseract.image_to_string(fd.askopenfilename(), lang='tam')
    print(ocr)
    doc = docx.Document()
    doc.add_paragraph(ocr)
    doc.save(fd.asksaveasfilename(defaultextension='.docx'))

def upscale():
    new = 1
    url = 'https://colab.research.google.com/github/tensorflow/hub/blob/master/examples/colab/image_enhancing.ipynb'
    webbrowser.open(url,new=new)


#to jpg
button_image_0 = ImageTk.PhotoImage(Image.open('D:\pyexpo\c1.jpg'))
button_0 = Button(
    image=button_image_0,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:to_jpg()
    ,

    relief="flat"
)

button_0.place(
    x=535.8,
    y=221.4,
    width=79,
    height=27
)

#to pdf
button_image_1 = ImageTk.PhotoImage(Image.open('D:\pyexpo\c1.jpg'))
button_1 = Button(
    image=button_image_1,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:to_pdf()
    ,

    relief="flat"
)

button_1.place(
    x=846.3,
    y=221.4,
    width=79,
    height=27
)

#to tiff
button_image_2 = ImageTk.PhotoImage(Image.open('D:\pyexpo\c1.jpg'))
button_2 = Button(
    image=button_image_2,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:to_tiff()
    ,

    relief="flat"
)

button_2.place(
    x=1150.3,
    y=221.4,
    width=79,
    height=27
)

#to docxe
button_image_3 = ImageTk.PhotoImage(Image.open('D:\pyexpo\e1.jpg'))
button_3 = Button(
    image=button_image_3,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:eng()
    ,

    relief="flat"
)

button_3.place(
    x=769.7,
    y=474.9,
    width=74,
    height=25
)

#to docxt
button_image_4 = ImageTk.PhotoImage(Image.open('D:\pyexpo\sr2.jpg'))
button_4 = Button(
    image=button_image_4,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:tam()
    ,

    relief="flat"
)

button_4.place(
    x=850.9,
    y=474.8,
    width=69,
    height=25
)

#to gif
button_image_5 = ImageTk.PhotoImage(Image.open('D:\pyexpo\c1.jpg'))
button_5 = Button(
    image=button_image_5,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:to_gif()
    ,

    relief="flat"
)

button_5.place(
    x=1150.3,
    y=474.8,
    width=79,
    height=27
)


#to png
button_image_7 = ImageTk.PhotoImage(Image.open('D:\pyexpo\c1.jpg'))
button_7 = Button(
    image=button_image_7,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:to_png()
    ,

    relief="flat"
)

button_7.place(
    x=535.8,
    y=474.8,
    width=87,
    height=29
)

#up
button_image_8 = ImageTk.PhotoImage(Image.open('D:\pyexpo\\ui.jpg'))
button_8 = Button(
    image=button_image_8,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:upscale()
    ,

    relief="flat"
)

button_8.place(
    x=535.8,
    y=724.7,
    width=80,
    height=27
)

#back
button_image_9 = ImageTk.PhotoImage(Image.open('D:\pyexpo\pg2bk.jpg'))
button_9 = Button(
    image=button_image_9,
    borderwidth=0,
    highlightthickness=0,
    command=lambda:back()
    ,

    relief="flat"
)

button_9.place(
    x=1115.3,
    y=753.1,
    width=102,
    height=48
)
root.mainloop()